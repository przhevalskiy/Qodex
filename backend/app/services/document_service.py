from typing import List, Optional, Dict, Any, Tuple
from pathlib import Path
import tiktoken
import uuid
import logging
import re
import json
from pypdf import PdfReader
from docx import Document as DocxDocument
import io

from app.models.document import Document, DocumentChunk
from app.services.pinecone_service import get_pinecone_service
from app.utils.course_utils import extract_course_title_from_content

logger = logging.getLogger(__name__)

# Persist document metadata so list_documents() survives server restarts
_REGISTRY_DIR = Path(__file__).resolve().parent.parent.parent / "data"
_REGISTRY_PATH = _REGISTRY_DIR / "document_registry.json"


class DocumentService:
    """Service for processing and managing documents."""

    def __init__(self):
        self.pinecone = get_pinecone_service()
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
        self.max_chunk_tokens = 500
        self.chunk_overlap = 50
        # In-memory cache — hydrated from disk on startup
        self._documents: Dict[str, Document] = {}
        self._load_registry()
        # Instructor → document_ids index for entity-first retrieval
        self.instructor_index: Dict[str, List[str]] = {}
        self._build_instructor_index()

    # =========================================================================
    # Document registry persistence
    # =========================================================================

    def _load_registry(self) -> None:
        """Load document metadata from disk (survives server restarts)."""
        if not _REGISTRY_PATH.exists():
            return
        try:
            data = json.loads(_REGISTRY_PATH.read_text())
            for entry in data:
                doc = Document.model_validate(entry)
                self._documents[doc.id] = doc
            logger.info(f"Loaded {len(self._documents)} documents from registry")
        except Exception as e:
            logger.warning(f"Failed to load document registry: {e}")

    def _save_registry(self) -> None:
        """Persist document metadata to disk."""
        try:
            _REGISTRY_DIR.mkdir(parents=True, exist_ok=True)
            data = [doc.model_dump(mode="json") for doc in self._documents.values()]
            _REGISTRY_PATH.write_text(json.dumps(data, default=str))
        except Exception as e:
            logger.warning(f"Failed to save document registry: {e}")

    # =========================================================================
    # Instructor index for entity-first retrieval
    # =========================================================================

    def _extract_instructor_from_filename(self, filename: str) -> Optional[str]:
        """Extract instructor name from filename using common patterns.

        Patterns:
        - "InstructorName_Topic.pdf" (at beginning)
        - "SyllabusReadings_Topic_InstructorName_Term.pdf"
        - "CourseCode_InstructorName.pdf"
        - "InstructorName_Topic.pdf" (at end)
        """
        # Remove extension
        name = re.sub(r'\.[^.]+$', '', filename)

        # Pattern 0: InstructorName at beginning (e.g., "BruceUsher_American-Innovation...")
        # Match CamelCase or single capitalized words at the start
        match = re.match(r'^([A-Z][a-z]+(?:[A-Z][a-z]+)*)_', name)
        if match:
            return match.group(1)

        # Pattern 1: ..._InstructorName_Fall2025 or ..._InstructorName_Spring2024
        match = re.search(r'_([A-Z][a-z]+(?:[A-Z][a-z]+)?)_(?:Fall|Spring|Summer)\d{4}', name)
        if match:
            return match.group(1)

        # Pattern 2: CourseCode_InstructorName (e.g., "SUMA PS5021_Kliegman")
        match = re.search(r'[A-Z]{3,5}\s+[A-Z]{2}\d{4}_([A-Z][a-z]+)', name)
        if match:
            return match.group(1)

        # Pattern 3: InstructorName at end before extension (e.g., "Topic-InstructorName.pdf")
        parts = re.split(r'[-_]', name)
        if len(parts) >= 2:
            last_part = parts[-1]
            if last_part and last_part[0].isupper() and not last_part.isupper():
                # Looks like a capitalized name (not an acronym)
                return last_part

        return None

    def _build_instructor_index(self) -> None:
        """Build reverse index: instructor name → list of document IDs."""
        self.instructor_index.clear()
        for doc in self._documents.values():
            instructor = self._extract_instructor_from_filename(doc.filename)
            if instructor:
                # Normalize: "TKhotin" → "tkhotin", "BruceUsher" → "bruce usher"
                normalized = re.sub(r'(?<!^)(?=[A-Z])', ' ', instructor).lower()
                if normalized not in self.instructor_index:
                    self.instructor_index[normalized] = []
                self.instructor_index[normalized].append(doc.id)

        if self.instructor_index:
            logger.info(f"Built instructor index: {len(self.instructor_index)} instructors, "
                       f"{sum(len(docs) for docs in self.instructor_index.values())} documents")

    def get_documents_by_instructor(self, instructor_name: str) -> Optional[List[str]]:
        """Get document IDs for a given instructor name (case-insensitive, exact match only).

        Returns:
            List of document IDs if found, None otherwise.
        """
        normalized = instructor_name.lower().strip()

        # Exact match only (no fuzzy matching to prevent cross-contamination)
        if normalized in self.instructor_index:
            return self.instructor_index[normalized]

        return None

    async def bootstrap_registry(self) -> int:
        """Discover all documents from Pinecone and populate the registry.

        This is needed after a server restart when the registry JSON doesn't
        exist yet (i.e. all documents were uploaded before persistence was added).

        Returns the number of new documents discovered.
        """
        logger.info("Bootstrapping document registry from Pinecone...")

        # Step 1: List all vector IDs
        all_ids = await self.pinecone.list_all_vector_ids()
        if not all_ids:
            logger.info("No vectors found in Pinecone — nothing to bootstrap")
            return 0

        # Step 2: Group by document_id (vector ID format: {uuid}_{chunk_index})
        doc_chunks: dict[str, list[str]] = {}
        for vid in all_ids:
            parts = vid.rsplit("_", 1)
            if len(parts) == 2 and parts[1].isdigit():
                doc_id = parts[0]
            else:
                doc_id = vid  # fallback: treat full ID as document
            doc_chunks.setdefault(doc_id, []).append(vid)

        # Step 3: Skip documents we already know about
        new_doc_ids = [did for did in doc_chunks if did not in self._documents]
        if not new_doc_ids:
            logger.info("Registry already contains all Pinecone documents")
            return 0

        # Step 4: Fetch one representative vector per new document to get metadata
        discovered = 0
        # Batch fetch in groups of 100 (Pinecone limit)
        representative_ids = [doc_chunks[did][0] for did in new_doc_ids]
        for i in range(0, len(representative_ids), 100):
            batch = representative_ids[i : i + 100]
            try:
                fetched = await self.pinecone.fetch_vectors(batch)
            except Exception as e:
                logger.warning(f"Failed to fetch vector batch: {e}")
                continue

            for vid, vec_data in fetched.items():
                metadata = vec_data.get("metadata", {})
                # Recover the document_id from the vector ID
                parts = vid.rsplit("_", 1)
                doc_id = parts[0] if (len(parts) == 2 and parts[1].isdigit()) else vid

                if doc_id in self._documents:
                    continue  # already known

                chunk_ids = sorted(doc_chunks.get(doc_id, []))
                doc = Document(
                    id=doc_id,
                    filename=metadata.get("filename", "Unknown Document"),
                    content_type=metadata.get("content_type", "application/pdf"),
                    file_size=metadata.get("file_size", 0),
                    chunk_count=len(chunk_ids),
                    chunk_ids=chunk_ids,
                    is_embedded=True,
                )
                self._documents[doc_id] = doc
                discovered += 1

        if discovered > 0:
            self._save_registry()
            logger.info(f"Bootstrapped {discovered} documents from Pinecone")
        else:
            logger.info("No new documents discovered during bootstrap")

        return discovered

    # =========================================================================
    # Private helpers for Pinecone operations
    # =========================================================================

    async def _fetch_chunks_from_pinecone(self, document_id: str) -> List[Dict[str, Any]]:
        """
        Fetch and sort chunks for a document from Pinecone.

        Args:
            document_id: The document ID to fetch chunks for

        Returns:
            List of chunks sorted by chunk_index, or empty list if not found
        """
        try:
            chunks = await self.pinecone.get_chunks_by_document(document_id)
            if not chunks:
                return []

            # Sort by chunk_index for consistent ordering
            return sorted(
                chunks,
                key=lambda c: c.get("metadata", {}).get("chunk_index", 0)
            )
        except Exception as e:
            logger.error(f"Failed to fetch chunks from Pinecone for {document_id}: {e}")
            return []

    def _reconstruct_document_from_chunks(
        self,
        document_id: str,
        chunks: List[Dict[str, Any]]
    ) -> Optional[Document]:
        """
        Reconstruct a Document object from Pinecone chunks.

        Args:
            document_id: The document ID
            chunks: List of chunks (should already be sorted)

        Returns:
            Document object or None if chunks are empty
        """
        if not chunks:
            return None

        first_chunk = chunks[0]
        metadata = first_chunk.get("metadata", {})

        return Document(
            id=document_id,
            filename=metadata.get("filename", "Unknown Document"),
            content_type=metadata.get("content_type", "application/pdf"),
            file_size=metadata.get("file_size", 0),
            chunk_count=len(chunks),
            chunk_ids=[chunk["id"] for chunk in chunks]
        )

    def _count_tokens(self, text: str) -> int:
        """Count tokens in text."""
        return len(self.tokenizer.encode(text))

    def _chunk_text(self, text: str) -> List[Dict[str, Any]]:
        """
        Split text into chunks while preserving structure.

        Returns list of dicts with 'content' and 'type' (heading, paragraph, list).
        """
        # Split into paragraphs (preserve structure)
        paragraphs = self._split_into_paragraphs(text)

        chunks = []
        current_chunk_parts = []
        current_tokens = 0

        for para in paragraphs:
            para_text = para["content"]
            para_tokens = self._count_tokens(para_text)

            # If single paragraph exceeds limit, split by sentences
            if para_tokens > self.max_chunk_tokens:
                # Flush current chunk first
                if current_chunk_parts:
                    chunks.append(self._merge_chunk_parts(current_chunk_parts))
                    current_chunk_parts = []
                    current_tokens = 0

                # Split large paragraph into sentence-based chunks
                sentence_chunks = self._split_paragraph_by_sentences(para)
                chunks.extend(sentence_chunks)
            elif current_tokens + para_tokens > self.max_chunk_tokens:
                # Flush current chunk and start new one
                if current_chunk_parts:
                    chunks.append(self._merge_chunk_parts(current_chunk_parts))
                current_chunk_parts = [para]
                current_tokens = para_tokens
            else:
                # Add to current chunk
                current_chunk_parts.append(para)
                current_tokens += para_tokens

        # Flush remaining
        if current_chunk_parts:
            chunks.append(self._merge_chunk_parts(current_chunk_parts))

        return chunks

    def _split_into_paragraphs(self, text: str) -> List[Dict[str, Any]]:
        """Split text into paragraphs with type detection."""
        paragraphs = []
        # Split on double newlines or single newlines followed by patterns
        raw_paragraphs = text.split('\n\n')

        for raw in raw_paragraphs:
            raw = raw.strip()
            if not raw:
                continue

            # Further split on single newlines that indicate structure
            lines = raw.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                para_type = self._detect_paragraph_type(line)
                paragraphs.append({
                    "content": line,
                    "type": para_type
                })

        return paragraphs

    def _detect_paragraph_type(self, text: str) -> str:
        """Detect the type of a text block."""
        text_stripped = text.strip()

        # Heading patterns (short, often title case or all caps)
        if len(text_stripped) < 100:
            # All caps heading
            if text_stripped.isupper() and len(text_stripped.split()) <= 10:
                return "heading"
            # Numbered heading (e.g., "1. Introduction", "Chapter 2")
            if text_stripped[:2].replace('.', '').isdigit():
                return "heading"
            # Title case and short
            words = text_stripped.split()
            if len(words) <= 8 and sum(1 for w in words if w[0].isupper()) >= len(words) * 0.6:
                return "heading"

        # List item patterns
        if text_stripped.startswith(('•', '-', '*', '●', '○')):
            return "list_item"
        if len(text_stripped) > 2 and text_stripped[0].isdigit() and text_stripped[1] in '.):':
            return "list_item"

        return "paragraph"

    def _split_paragraph_by_sentences(self, para: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Split a large paragraph into sentence-based chunks."""
        text = para["content"]
        para_type = para["type"]

        # Simple sentence split (handles common cases)
        sentences = []
        current = ""
        for char in text:
            current += char
            if char in '.!?' and len(current) > 20:
                sentences.append(current.strip())
                current = ""
        if current.strip():
            sentences.append(current.strip())

        # Group sentences into chunks
        chunks = []
        current_chunk = []
        current_tokens = 0

        for sentence in sentences:
            sent_tokens = self._count_tokens(sentence)
            if current_tokens + sent_tokens > self.max_chunk_tokens:
                if current_chunk:
                    chunks.append({
                        "content": ' '.join(current_chunk),
                        "type": para_type
                    })
                current_chunk = [sentence]
                current_tokens = sent_tokens
            else:
                current_chunk.append(sentence)
                current_tokens += sent_tokens

        if current_chunk:
            chunks.append({
                "content": ' '.join(current_chunk),
                "type": para_type
            })

        return chunks

    def _merge_chunk_parts(self, parts: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Merge multiple paragraph parts into a single chunk."""
        if not parts:
            return {"content": "", "type": "paragraph"}

        # Join with double newlines to preserve structure
        content = '\n\n'.join(p["content"] for p in parts)

        # Determine dominant type
        types = [p["type"] for p in parts]
        if types[0] == "heading":
            chunk_type = "heading"
        elif "list_item" in types and types.count("list_item") > len(types) / 2:
            chunk_type = "list"
        else:
            chunk_type = "paragraph"

        return {"content": content, "type": chunk_type}

    def _extract_text_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF content with cleanup and normalization."""
        reader = PdfReader(io.BytesIO(content))
        text = ""
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n"

        # Clean up the extracted text
        text = self._clean_pdf_text(text)
        return text

    def _clean_pdf_text(self, text: str) -> str:
        """
        Clean up PDF extracted text by removing watermarks and
        rejoining fragmented lines.

        Uses only structural / statistical heuristics — no keyword matching.
        Default behaviour is JOIN; only starts a new line on strong signals.
        """
        # Phase 1: Remove noise
        text = re.sub(r'Downloaded from Qodex[^\n]*', '', text, flags=re.IGNORECASE)
        text = re.sub(r'^\s*Page\s+\d+\s+of\s+\d+\s*$', '', text, flags=re.MULTILINE)
        text = re.sub(r'^\s*\d{1,3}\s*$', '', text, flags=re.MULTILINE)

        # Phase 2: Defragment — default to joining
        lines = text.split('\n')
        cleaned_lines: list[str] = []
        buffer: list[str] = []

        def flush():
            if buffer:
                cleaned_lines.append(' '.join(buffer))
                buffer.clear()

        for line in lines:
            stripped = line.strip()

            if not stripped:
                flush()
                cleaned_lines.append('')
                continue

            if not buffer:
                buffer.append(stripped)
                continue

            if self._is_new_logical_line(stripped, buffer):
                flush()
                buffer.append(stripped)
            else:
                # Handle hyphenation
                if buffer and buffer[-1].endswith('-') and stripped[0:1].islower():
                    buffer[-1] = buffer[-1][:-1]
                buffer.append(stripped)

        flush()

        result = '\n'.join(cleaned_lines)
        result = re.sub(r'\n{3,}', '\n\n', result)
        result = re.sub(r'[ \t]{2,}', ' ', result)
        return result.strip()

    @staticmethod
    def _is_all_caps(text: str) -> bool:
        """Check if text is ALL-CAPS (every word uppercase, at least one 2+ char word)."""
        has_substantial = False
        for word in text.split():
            letters = re.sub(r'[^a-zA-Z]', '', word)
            if not letters:
                continue
            if letters != letters.upper():
                return False
            if len(letters) >= 2:
                has_substantial = True
        return has_substantial

    @staticmethod
    def _is_title_like(text: str) -> bool:
        """Check if text looks like a title: 2-8 words, mostly capitalised, no sentence-end punctuation."""
        words = text.split()
        wc = len(words)
        if wc < 2 or wc > 8:
            return False
        if len(text) >= 80:
            return False
        if re.search(r'[.!?]\s*$', text):
            return False
        if not text[0].isupper():
            return False
        cap_count = sum(1 for w in words if w[0].isupper())
        return cap_count >= wc * 0.5

    def _is_new_logical_line(self, line: str, buffer: list) -> bool:
        """
        Should this line start a new logical line rather than joining the buffer?

        AGGRESSIVE JOIN: only breaks for unambiguous syntax (bullets, numbered
        lists).  All heading/structure detection is deferred to the frontend
        rendering pipeline.
        """
        # Syntax-based structural elements — always start new
        if re.match(r'^[•●○]\s', line):
            return True
        if re.match(r'^[-*]\s', line) and len(line) < 200:
            return True
        if re.match(r'^\d+[.)]\s', line):
            return True
        if re.match(r'^#{1,3}\s', line):
            return True
        if re.match(r'^[-_*]{3,}\s*$', line):
            return True

        # After sentence boundary, only break if new line is substantial (>= 40 chars)
        prev_text = ' '.join(buffer)
        prev_ended_sentence = bool(re.search(r'[.!?]\s*$', prev_text))
        if prev_ended_sentence and line[0:1].isupper() and len(line) >= 40:
            return True

        # Default: JOIN
        return False

    def _extract_text_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX content."""
        doc = DocxDocument(io.BytesIO(content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text

    def _extract_text(self, content: bytes, content_type: str, filename: str) -> str:
        """Extract text from document based on type."""
        if content_type == "application/pdf" or filename.endswith(".pdf"):
            return self._extract_text_from_pdf(content)
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or filename.endswith(".docx"):
            return self._extract_text_from_docx(content)
        elif content_type.startswith("text/") or filename.endswith((".txt", ".md")):
            return content.decode("utf-8")
        else:
            raise ValueError(f"Unsupported content type: {content_type}")

    async def process_document(
        self,
        filename: str,
        content: bytes,
        content_type: str
    ) -> Document:
        """
        Process and embed a document.

        Args:
            filename: Name of the file
            content: Raw file content
            content_type: MIME type of the file

        Returns:
            Document object with metadata
        """
        # Create document record
        doc_id = str(uuid.uuid4())
        document = Document(
            id=doc_id,
            filename=filename,
            content_type=content_type,
            file_size=len(content)
        )

        # Extract text
        text = self._extract_text(content, content_type, filename)

        # Chunk the text (now returns structured chunks with type)
        chunks = self._chunk_text(text)
        document.chunk_count = len(chunks)

        # Create embeddings for chunk content
        chunk_contents = [c["content"] for c in chunks]
        embeddings = await self.pinecone.create_embeddings_batch(chunk_contents)

        # Extract course title from first chunk for metadata tagging.
        # Stored on every chunk so Pinecone metadata filters can scope
        # retrieval to a specific course at query time.
        course_name = extract_course_title_from_content(
            [c["content"] for c in chunks[:1]]
        ) or ""

        # Prepare vectors for Pinecone with structure metadata
        vectors = []
        chunk_ids = []
        for i, (chunk_data, embedding) in enumerate(zip(chunks, embeddings)):
            chunk_id = f"{doc_id}_{i}"
            chunk_ids.append(chunk_id)
            vectors.append({
                "id": chunk_id,
                "values": embedding,
                "metadata": {
                    "document_id": doc_id,
                    "filename": filename,
                    "chunk_index": i,
                    "content": chunk_data["content"],
                    "content_type": chunk_data["type"],  # heading, paragraph, list
                    "course_name": course_name,
                }
            })

        # Upsert to Pinecone
        await self.pinecone.upsert_vectors(vectors)

        # Update document record
        document.chunk_ids = chunk_ids
        document.is_embedded = True

        # Store document and persist registry
        self._documents[doc_id] = document
        self._save_registry()

        return document

    async def delete_document(self, document_id: str) -> bool:
        """
        Delete a document and its vectors.

        Args:
            document_id: ID of the document to delete

        Returns:
            True if deleted, False if not found
        """
        # Try cache first, then Pinecone
        document = self._documents.get(document_id)
        if not document:
            chunks = await self._fetch_chunks_from_pinecone(document_id)
            document = self._reconstruct_document_from_chunks(document_id, chunks)

        if not document:
            return False

        # Delete vectors from Pinecone
        await self.pinecone.delete_vectors(ids=document.chunk_ids)

        # Remove from cache and persist
        self._documents.pop(document_id, None)
        self._save_registry()
        return True

    async def get_document(self, document_id: str) -> Optional[Document]:
        """
        Get a document by ID.

        Checks in-memory cache first, falls back to Pinecone reconstruction.
        """
        # Check cache first
        if document := self._documents.get(document_id):
            return document

        # Reconstruct from Pinecone
        chunks = await self._fetch_chunks_from_pinecone(document_id)
        return self._reconstruct_document_from_chunks(document_id, chunks)

    def list_documents(self) -> List[Document]:
        """List all documents."""
        return list(self._documents.values())

    async def search_documents(
        self,
        query: str,
        top_k: int = 5,
        document_ids: Optional[List[str]] = None
    ) -> str:
        """
        Search documents and return formatted context.

        Args:
            query: Search query
            top_k: Number of chunks to retrieve
            document_ids: Optional filter by document IDs

        Returns:
            Formatted context string for the AI
        """
        results = await self.pinecone.search_documents(
            query=query,
            top_k=top_k,
            document_ids=document_ids
        )

        if not results:
            return ""

        context_parts = []
        for result in results:
            if result.get("metadata"):
                filename = result["metadata"].get("filename", "Unknown")
                content = result["metadata"].get("content", "")
                context_parts.append(f"[From {filename}]:\n{content}")

        return "\n\n---\n\n".join(context_parts)

    async def get_document_content(self, document_id: str) -> Dict[str, Any]:
        """
        Get full document content for preview.

        Args:
            document_id: ID of document

        Returns:
            Dictionary with document metadata and full content

        Raises:
            ValueError: If document not found
        """
        # Fetch chunks once - used for both document reconstruction and content
        chunks = await self._fetch_chunks_from_pinecone(document_id)
        if not chunks:
            raise ValueError(f"Document not found: {document_id}")

        # Get or reconstruct document
        document = self._documents.get(document_id)
        if not document:
            document = self._reconstruct_document_from_chunks(document_id, chunks)

        # Build content from chunks (already sorted by _fetch_chunks_from_pinecone)
        chunk_contents = []
        content_parts = []

        for chunk in chunks:
            metadata = chunk.get("metadata", {})
            content = metadata.get("content", "")
            if content:
                chunk_contents.append({
                    "id": chunk["id"],
                    "content": content,
                    "chunk_index": metadata.get("chunk_index", 0),
                    "content_type": metadata.get("content_type", "paragraph")  # heading, paragraph, list
                })
                content_parts.append(content)

        return {
            "id": document.id,
            "filename": document.filename,
            "content_type": document.content_type,
            "file_size": document.file_size,
            "chunk_count": document.chunk_count,
            "full_content": "\n\n".join(content_parts),
            "chunks": chunk_contents
        }
    
    async def get_document_chunks(self, document_id: str) -> List[Dict[str, Any]]:
        """
        Get document chunks for preview.

        Args:
            document_id: ID of the document

        Returns:
            List of chunk metadata (id, chunk_index, filename)

        Raises:
            ValueError: If document not found
        """
        chunks = await self._fetch_chunks_from_pinecone(document_id)
        if not chunks:
            raise ValueError(f"Document not found: {document_id}")

        # Extract filename from first chunk
        default_filename = chunks[0].get("metadata", {}).get("filename", "Unknown Document")

        # Build response (chunks already sorted by _fetch_chunks_from_pinecone)
        return [
            {
                "id": chunk["id"],
                "chunk_index": chunk.get("metadata", {}).get("chunk_index", 0),
                "filename": chunk.get("metadata", {}).get("filename", default_filename),
                "content_type": chunk.get("metadata", {}).get("content_type", "paragraph")
            }
            for chunk in chunks
            if chunk.get("metadata")
        ]


# Singleton instance
_document_service: Optional[DocumentService] = None

def get_document_service() -> DocumentService:
    """Get the singleton DocumentService instance."""
    global _document_service
    if _document_service is None:
        _document_service = DocumentService()
    return _document_service
